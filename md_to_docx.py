"""Convert the Z Drive Storage Audit Report from Markdown to Word (.docx)."""

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

INPUT  = r"c:\Users\SAguiar\Downloads\extractemails.py\Z_Drive_Storage_Audit_Report.md"
OUTPUT = r"c:\Users\SAguiar\Downloads\Z_Drive_Storage_Audit_Report.docx"


def set_cell_text(cell, text, bold=False, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold


def add_table(doc, header_row, data_rows):
    cols = len(header_row)
    table = doc.add_table(rows=1, cols=cols, style="Light Grid Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(header_row):
        set_cell_text(table.rows[0].cells[i], h.strip(), bold=True, size=9)
    for row_data in data_rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_data):
            if i < cols:
                set_cell_text(cells[i], val.strip(), size=9)
    doc.add_paragraph()  # spacing


def parse_table_block(lines):
    """Parse markdown table lines into header + data rows."""
    header = [c.strip() for c in lines[0].strip("|").split("|")]
    data = []
    for line in lines[2:]:  # skip separator
        row = [c.strip() for c in line.strip("|").split("|")]
        data.append(row)
    return header, data


def main():
    with open(INPUT, "r", encoding="utf-8") as f:
        raw = f.read()

    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = raw.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # --- Tables ---
        if "|" in line and i + 1 < len(lines) and re.match(r"^\s*\|[\s\-:|]+\|\s*$", lines[i + 1]):
            table_lines = []
            while i < len(lines) and "|" in lines[i]:
                table_lines.append(lines[i])
                i += 1
            header, data = parse_table_block(table_lines)
            add_table(doc, header, data)
            continue

        # --- Headings ---
        if line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=0)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue

        # --- Horizontal rule ---
        if re.match(r"^-{3,}$", line.strip()):
            doc.add_paragraph("_" * 60)
            i += 1
            continue

        # --- Blank line ---
        if line.strip() == "":
            i += 1
            continue

        # --- Bullet / numbered list ---
        bullet_match = re.match(r"^(\s*)[-*]\s+(.*)", line)
        num_match = re.match(r"^(\s*)\d+\.\s+(.*)", line)
        if bullet_match:
            indent = len(bullet_match.group(1)) // 2
            text = bullet_match.group(2)
            p = doc.add_paragraph(style="List Bullet")
            # Handle bold within text
            parts = re.split(r"(\*\*.*?\*\*)", text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
            if indent > 0:
                p.paragraph_format.left_indent = Inches(0.25 * indent)
            i += 1
            continue
        if num_match:
            indent = len(num_match.group(1)) // 2
            text = num_match.group(2)
            p = doc.add_paragraph(style="List Number")
            parts = re.split(r"(\*\*.*?\*\*)", text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
            if indent > 0:
                p.paragraph_format.left_indent = Inches(0.25 * indent)
            i += 1
            continue

        # --- Regular paragraph ---
        p = doc.add_paragraph()
        # Handle bold and inline code
        parts = re.split(r"(\*\*.*?\*\*|`[^`]+`)", line)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("`") and part.endswith("`"):
                run = p.add_run(part[1:-1])
                run.font.name = "Consolas"
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x80, 0x00, 0x00)
            else:
                p.add_run(part)
        i += 1

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    main()
